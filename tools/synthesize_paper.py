from datetime import datetime
from docx import Document
import os


OUTPUT_DIR = "results"
OUTPUT_FILE = "review_paper.docx"


def synthesize_paper(repo, workflow_id, execution_attempt_id=None, **kwargs):

    rows = repo.fetch_all(
        """
        SELECT section_name, content
        FROM DraftSection
        WHERE workflow_id = ?
        """,
        (workflow_id,)
    )

    if not rows:
        return {
            "status": "error",
            "data": None,
            "error": "No draft sections found."
        }

    # store sections
    paper = {}

    for r in rows:
        paper[r["section_name"]] = r["content"]

    # ensure results folder exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()
    doc.add_heading("Automated Literature Review", 0)

    # write sections
    for section, text in paper.items():

        doc.add_heading(section, level=1)
        doc.add_paragraph(text)

    output_path = os.path.join(OUTPUT_DIR, OUTPUT_FILE)

    doc.save(output_path)

    # store final sections in DB
    with repo.transaction() as cursor:

        for section, text in paper.items():

            cursor.execute(
                """
                INSERT INTO FinalPaperSection (
                    workflow_id,
                    section_name,
                    content
                )
                VALUES (?, ?, ?)
                """,
                (
                    workflow_id,
                    section,
                    text
                )
            )

    return {
        "status": "success",
        "data": output_path,
        "error": None
    }